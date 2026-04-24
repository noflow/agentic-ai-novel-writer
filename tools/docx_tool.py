"""
Word Document tool -- creates formatted .docx files.

Install: pip install python-docx
"""

import os
import re
from pathlib import Path
from tools.base import BaseTool

# Try importing python-docx at module level
_DOCX_AVAILABLE = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    pass


def _get_output_dir() -> Path:
    out = Path(__file__).parent.parent / "output"
    out.mkdir(exist_ok=True)
    return out


def _resolve_path(filepath: str) -> str:
    filepath = filepath.strip().strip("\"'")
    p = Path(filepath)
    if p.parent == Path(".") or str(p.parent) == ".":
        return str(_get_output_dir() / p.name)
    return str(p.resolve())


def _add_formatted_text(paragraph, text: str):
    """Parse inline **bold** and *italic* markers and add styled runs."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(12)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Georgia"
            run.font.size = Pt(12)
        else:
            run = paragraph.add_run(part)
            run.font.name = "Georgia"
            run.font.size = Pt(12)


def _build_docx(content: str, filepath: str) -> str:
    """Build a .docx file from formatted text content."""
    if not _DOCX_AVAILABLE:
        return (
            "Error: python-docx not installed. Run:\n"
            "  pip install python-docx"
        )

    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(12)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Parse and add content line by line
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Page break
        if line.strip() == "---":
            doc.add_page_break()
            i += 1
            continue

        # Title (# )
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = "Georgia"
            p.paragraph_format.space_after = Pt(24)
            i += 1
            continue

        # Chapter heading (## )
        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Georgia"
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(18)
            i += 1
            continue

        # Section heading (### )
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Empty line
        if line.strip() == "":
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        _add_formatted_text(p, line)
        i += 1

    # Save
    resolved = _resolve_path(filepath)
    parent = os.path.dirname(resolved)
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(resolved)

    return (
        f"Successfully created Word document: {resolved}\n"
        f"Pages: ~{len(lines) // 30 + 1} (estimated)\n"
        f"You can open this in Microsoft Word, Google Docs, or LibreOffice."
    )


class CreateDocxTool(BaseTool):
    name = "create_docx"
    description = (
        "Create a formatted Word document (.docx) from text content. "
        "The content can include formatting markers:\n"
        "  # Title (large, bold, centered)\n"
        "  ## Chapter heading (bold, centered)\n"
        "  ### Section heading (bold)\n"
        "  --- (page break)\n"
        "  *italic* and **bold**\n"
        "  Regular text becomes indented paragraphs.\n"
        "You can also set source_file to convert a .txt file."
    )
    input_schema = {
        "type": "object",
        "properties": {
            "filepath": {
                "type": "string",
                "description": "Output filename ending in .docx",
            },
            "content": {
                "type": "string",
                "description": "Text with formatting markers, OR leave empty and use source_file",
            },
            "source_file": {
                "type": "string",
                "description": "Optional: .txt file path to convert instead of using content",
            },
        },
        "required": ["filepath"],
    }

    def run(self, filepath: str = "", content: str = "", source_file: str = "", **kwargs) -> str:
        if not filepath:
            return "Error: No filepath provided."

        if not filepath.lower().endswith(".docx"):
            filepath += ".docx"

        # Read source file if given
        if source_file and not content:
            try:
                src = _resolve_path(source_file)
                with open(src, "r", encoding="utf-8") as f:
                    content = f.read()
            except Exception as e:
                return f"Error reading source file: {e}"

        if not content:
            return "Error: No content. Pass content or set source_file."

        try:
            return _build_docx(content, filepath)
        except Exception as e:
            return f"Error creating Word document: {e}"


class TxtToDocxTool(BaseTool):
    name = "convert_to_docx"
    description = (
        "Convert a .txt file to a formatted Word document (.docx). "
        "Automatically applies formatting for headings, paragraphs, "
        "and dialogue."
    )
    input_schema = {
        "type": "object",
        "properties": {
            "source_file": {
                "type": "string",
                "description": "Path to the .txt file to convert",
            },
            "output_file": {
                "type": "string",
                "description": "Output .docx filename (default: same name with .docx)",
            },
        },
        "required": ["source_file"],
    }

    def run(self, source_file: str = "", output_file: str = "", **kwargs) -> str:
        if not source_file:
            return "Error: No source_file provided."

        src = _resolve_path(source_file)
        if not os.path.exists(src):
            return f"Error: File not found: {source_file}"

        if not output_file:
            output_file = Path(src).stem + ".docx"
        if not output_file.lower().endswith(".docx"):
            output_file += ".docx"

        try:
            with open(src, "r", encoding="utf-8") as f:
                content = f.read()
            return _build_docx(content, output_file)
        except Exception as e:
            return f"Error converting to docx: {e}"
